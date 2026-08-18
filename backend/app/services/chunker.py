"""文档解析与切块模块。

支持 .docx / .pdf / .txt / .md 格式。
切块策略对法律文档做了专门优化：
- 优先识别「第X条」「第X章」等结构标记，以条款为切块边界；
- 保留每个块的元数据（来源文件、分类、文件名、页码、条款号）。
"""
from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

ARTICLE_RE = re.compile(r"^\s*(第[一二三四五六七八九十百千0-9]+[条章])\s*(.*)$")
LINE_BREAK_RE = re.compile(r"[ \t\u3000]+")


@dataclass
class Chunk:
    """知识库中的一个检索单元。"""

    text: str
    doc_id: str              # 文档唯一 ID（相对路径）
    title: str               # 文件名（不含扩展名）
    category: str            # 一级分类目录名
    article: str = ""        # 条款号，如「第三十九条」
    page: int = 0            # 页码（PDF）
    chunk_index: int = 0     # 块序号
    meta: dict = field(default_factory=dict)


@dataclass
class ParsedDoc:
    """解析结果：段落列表 + 页面信息。"""

    doc_id: str
    title: str
    category: str
    paragraphs: list[str] = field(default_factory=list)
    page_of_para: list[int] = field(default_factory=list)  # 每段所在页码
    meta: dict = field(default_factory=dict)


def detect_category(path: Path, kb_root: Path) -> str:
    """根据一级目录名推断分类，如 宪法/法律/行政法规/司法解释。"""
    try:
        rel = path.relative_to(kb_root)
        parts = rel.parts
        if len(parts) >= 2:
            return parts[0]
        return "未分类"
    except ValueError:
        return "未分类"


def _norm(text: str) -> str:
    return LINE_BREAK_RE.sub(" ", text).strip()


def parse_docx(path: Path, kb_root: Path) -> ParsedDoc:
    import docx  # python-docx

    doc = docx.Document(str(path))
    paras: list[str] = []
    for p in doc.paragraphs:
        t = _norm(p.text)
        if t:
            paras.append(t)
    # 表格内容也提取（法规常有表格）
    for table in doc.tables:
        for row in table.rows:
            cells = [_norm(c.text) for c in row.cells if _norm(c.text)]
            if cells:
                paras.append(" | ".join(cells))
    rel = path.relative_to(kb_root) if path.is_relative_to(kb_root) else path
    return ParsedDoc(
        doc_id=rel.as_posix(),
        title=path.stem,
        category=detect_category(path, kb_root),
        paragraphs=paras,
        page_of_para=[0] * len(paras),
    )


def parse_pdf(path: Path, kb_root: Path, use_ocr: bool = True) -> ParsedDoc:
    """解析 PDF。

    优先提取文本层；若某页无文本（扫描件），且启用 OCR，则自动用 RapidOCR 识别。
    """
    import pdfplumber

    from app.core.config import settings
    from app.services.ocr import ocr_pdf_page

    paras: list[str] = []
    pages: list[int] = []
    ocr_pages = 0
    with pdfplumber.open(str(path)) as pdf:
        for pno, page in enumerate(pdf.pages, start=1):
            text = page.extract_text() or ""
            lines = [_norm(l) for l in text.splitlines() if _norm(l)]
            if not lines and use_ocr and settings.ocr_enabled and not page.images == []:
                # 扫描页：走 OCR
                ocr_lines = ocr_pdf_page(page)
                lines = [_norm(l) for l in ocr_lines if _norm(l)]
                if lines:
                    ocr_pages += 1
            for line in lines:
                paras.append(line)
                pages.append(pno)
    rel = path.relative_to(kb_root) if path.is_relative_to(kb_root) else path
    return ParsedDoc(
        doc_id=rel.as_posix(),
        title=path.stem,
        category=detect_category(path, kb_root),
        paragraphs=paras,
        page_of_para=pages,
        meta={"ocr_pages": ocr_pages},
    )


def parse_txt(path: Path, kb_root: Path) -> ParsedDoc:
    content = path.read_text(encoding="utf-8", errors="ignore")
    paras = [_norm(p) for p in content.splitlines() if _norm(p)]
    rel = path.relative_to(kb_root) if path.is_relative_to(kb_root) else path
    return ParsedDoc(
        doc_id=rel.as_posix(),
        title=path.stem,
        category=detect_category(path, kb_root),
        paragraphs=paras,
        page_of_para=[0] * len(paras),
    )


def parse_document(path: Path, kb_root: Path) -> ParsedDoc:
    """按扩展名分发解析。"""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return parse_docx(path, kb_root)
    if suffix == ".pdf":
        return parse_pdf(path, kb_root)
    if suffix in (".txt", ".md", ".markdown"):
        return parse_txt(path, kb_root)
    raise ValueError(f"不支持的格式: {suffix}")


def _split_articles(paras: list[str], pages: list[int]) -> list[tuple[str, str, int]]:
    """将段落按「第X条」边界分组，返回 (文本, 条款号, 页码)。"""
    groups: list[tuple[list[str], str, int]] = []
    cur: list[str] = []
    cur_article = ""
    cur_page = 0

    def flush() -> None:
        nonlocal cur
        if cur:
            groups.append(("".join(cur), cur_article, cur_page))
            cur = []

    for para, page in zip(paras, pages):
        m = ARTICLE_RE.match(para)
        if m and (len(para) <= 60 or m.group(1) in ("第一章", "第二章", "第三章", "第四章", "第五章", "第六章", "第七章", "第八章", "第九章", "第十章", "第十一章", "第十二章", "第十三章", "第十四章", "第十五章", "第十六章")):
            flush()
            cur_article = m.group(1)
            cur_page = page
            cur.append(para)
        elif m and cur_article:
            # 连续多条「第X条」开头，视为新条款
            flush()
            cur_article = m.group(1)
            cur_page = page
            cur.append(para)
        else:
            cur.append(para)
    flush()
    return [(t, a, p) for t, a, p in groups if len(t) >= 5]


def chunk_document(doc: ParsedDoc, chunk_size: int = 600, overlap: int = 80) -> list[Chunk]:
    """把解析后的文档切成检索块。

    优先按条款聚合；条款过长时按 chunk_size 硬切并保留 overlap。
    """
    chunks: list[Chunk] = []
    idx = 0
    for text, article, page in _split_articles(doc.paragraphs, doc.page_of_para):
        text = text.strip()
        if not text:
            continue
        if len(text) <= chunk_size * 1.3:
            chunks.append(
                Chunk(
                    text=text,
                    doc_id=doc.doc_id,
                    title=doc.title,
                    category=doc.category,
                    article=article,
                    page=page,
                    chunk_index=idx,
                )
            )
            idx += 1
        else:
            # 长文本按字符窗口切
            step = chunk_size - overlap
            start = 0
            while start < len(text):
                seg = text[start : start + chunk_size]
                chunks.append(
                    Chunk(
                        text=seg,
                        doc_id=doc.doc_id,
                        title=doc.title,
                        category=doc.category,
                        article=article,
                        page=page,
                        chunk_index=idx,
                    )
                )
                idx += 1
                start += step
                if len(seg) < chunk_size:
                    break
    return chunks
